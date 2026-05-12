"""Render Markdown answers to polished PDF and DOCX files.

The chatbot used to dump the raw markdown source into the export. That meant
``### Moyenne S1`` and ``| UE | BCC | ... |`` showed up as literal characters
instead of a proper heading and a bordered table. This module replaces that
with a real renderer that:

* converts the markdown source to HTML with the official ``markdown`` package
  (tables, fenced code, smart lists);
* walks the resulting HTML tree with BeautifulSoup;
* emits ReportLab Platypus flowables (or python-docx blocks) with consistent
  styling, alternating-row tables, code monospace, blockquote bars, etc.

Both ``render_pdf`` and ``render_docx`` return ``bytes`` so the Streamlit
download button can serve them directly.
"""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Iterable

import markdown as md
from bs4 import BeautifulSoup, NavigableString, Tag

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Shared helpers
# --------------------------------------------------------------------------- #

PRIMARY = colors.HexColor("#2563eb")
INK = colors.HexColor("#0f172a")
MUTED = colors.HexColor("#475569")
RULE = colors.HexColor("#e2e8f0")
TABLE_HEADER_BG = colors.HexColor("#f1f5f9")
TABLE_ALT_BG = colors.HexColor("#fafafb")
CODE_BG = colors.HexColor("#f8fafc")
QUOTE_BAR = colors.HexColor("#bfdbfe")

MARKDOWN_EXTENSIONS = ["tables", "fenced_code", "sane_lists", "nl2br"]


def _markdown_to_soup(body: str) -> BeautifulSoup:
    html = md.markdown(body or "", extensions=MARKDOWN_EXTENSIONS, output_format="html5")
    return BeautifulSoup(html, "html.parser")


def _strip(text: str) -> str:
    """Collapse newlines/extra whitespace inside an HTML inline run."""
    return re.sub(r"\s+", " ", text or "").strip()


# --------------------------------------------------------------------------- #
# PDF renderer
# --------------------------------------------------------------------------- #

_FONTS_REGISTERED = False


def _register_fonts() -> tuple[str, str, str, str]:
    """Register Bitstream Vera (Unicode-capable) and return its font names."""
    global _FONTS_REGISTERED
    family = "BodyFont"
    bold = f"{family}-Bold"
    italic = f"{family}-Italic"
    bold_italic = f"{family}-BoldItalic"
    if _FONTS_REGISTERED:
        return family, bold, italic, bold_italic

    fonts_dir = Path(pdfmetrics.__file__).resolve().parent.parent / "fonts"
    pdfmetrics.registerFont(TTFont(family, str(fonts_dir / "Vera.ttf")))
    pdfmetrics.registerFont(TTFont(bold, str(fonts_dir / "VeraBd.ttf")))
    pdfmetrics.registerFont(TTFont(italic, str(fonts_dir / "VeraIt.ttf")))
    pdfmetrics.registerFont(TTFont(bold_italic, str(fonts_dir / "VeraBI.ttf")))
    pdfmetrics.registerFontFamily(
        family, normal=family, bold=bold, italic=italic, boldItalic=bold_italic
    )
    _FONTS_REGISTERED = True
    return family, bold, italic, bold_italic


def _build_styles() -> dict[str, ParagraphStyle]:
    body_font, bold_font, italic_font, _ = _register_fonts()
    base = getSampleStyleSheet()["BodyText"]
    common = dict(fontName=body_font, leading=15, textColor=INK)

    title = ParagraphStyle(
        "DocTitle",
        parent=base,
        fontName=bold_font,
        fontSize=22,
        leading=26,
        textColor=INK,
        spaceAfter=4,
    )
    subtitle = ParagraphStyle(
        "DocSubtitle",
        parent=base,
        fontName=italic_font,
        fontSize=10,
        textColor=MUTED,
        spaceAfter=14,
    )
    h1 = ParagraphStyle("H1", parent=base, fontName=bold_font, fontSize=18, leading=22, textColor=PRIMARY, spaceBefore=14, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=base, fontName=bold_font, fontSize=15, leading=19, textColor=INK, spaceBefore=12, spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=base, fontName=bold_font, fontSize=12.5, leading=16, textColor=INK, spaceBefore=10, spaceAfter=5)
    h4 = ParagraphStyle("H4", parent=base, fontName=bold_font, fontSize=11, leading=14, textColor=MUTED, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("Body", parent=base, fontSize=10.5, alignment=TA_LEFT, spaceAfter=6, **common)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=14, bulletIndent=2, spaceAfter=2)
    numbered = ParagraphStyle("Numbered", parent=body, leftIndent=18, bulletIndent=2, spaceAfter=2)
    quote = ParagraphStyle(
        "Quote",
        parent=body,
        leftIndent=14,
        textColor=MUTED,
        fontName=italic_font,
        spaceBefore=4,
        spaceAfter=8,
        borderPadding=4,
    )
    code = ParagraphStyle(
        "Code",
        parent=base,
        fontName="Courier",
        fontSize=9.5,
        leading=12,
        textColor=INK,
        backColor=CODE_BG,
        borderPadding=6,
        spaceBefore=4,
        spaceAfter=8,
    )
    table_cell = ParagraphStyle(
        "TableCell", parent=body, fontSize=9.5, leading=12, spaceAfter=0
    )
    table_header = ParagraphStyle(
        "TableHeader", parent=table_cell, fontName=bold_font, textColor=INK
    )
    return {
        "title": title,
        "subtitle": subtitle,
        "h1": h1,
        "h2": h2,
        "h3": h3,
        "h4": h4,
        "body": body,
        "bullet": bullet,
        "numbered": numbered,
        "quote": quote,
        "code": code,
        "cell": table_cell,
        "th": table_header,
    }


def _inline_to_paragraph_html(node: Tag | NavigableString) -> str:
    """Convert an inline HTML subtree to ReportLab-flavoured XML."""
    if isinstance(node, NavigableString):
        text = str(node)
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
    if not isinstance(node, Tag):
        return ""
    children = "".join(_inline_to_paragraph_html(child) for child in node.children)
    name = node.name.lower()
    if name in {"strong", "b"}:
        return f"<b>{children}</b>"
    if name in {"em", "i"}:
        return f"<i>{children}</i>"
    if name == "u":
        return f"<u>{children}</u>"
    if name in {"del", "s", "strike"}:
        return f'<font color="#9ca3af"><strike>{children}</strike></font>'
    if name == "code":
        return f'<font face="Courier" backColor="#eef2ff">{children}</font>'
    if name == "br":
        return "<br/>"
    if name == "a":
        href = node.get("href", "")
        href_safe = href.replace('"', "%22")
        if href:
            return f'<link href="{href_safe}" color="#2563eb">{children}</link>'
        return children
    if name == "span":
        return children
    return children


def _table_cell(cell_node: Tag, style: ParagraphStyle) -> Paragraph:
    body_html = "".join(_inline_to_paragraph_html(child) for child in cell_node.children)
    if not body_html.strip():
        body_html = "&nbsp;"
    return Paragraph(body_html, style)


def _build_table(table_node: Tag, styles: dict[str, ParagraphStyle]) -> Table:
    head_rows: list[list[Paragraph]] = []
    body_rows: list[list[Paragraph]] = []
    for tr in table_node.find_all("tr"):
        head_cells = tr.find_all("th")
        if head_cells:
            head_rows.append([_table_cell(th, styles["th"]) for th in head_cells])
        else:
            body_rows.append([_table_cell(td, styles["cell"]) for td in tr.find_all("td")])
    rows = head_rows + body_rows
    if not rows:
        return Table([[Paragraph("", styles["cell"])]])

    column_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < column_count:
            r.append(Paragraph("&nbsp;", styles["cell"]))

    table = Table(rows, repeatRows=len(head_rows), hAlign="LEFT")
    style_cmds = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LINEBELOW", (0, 0), (-1, -1), 0.4, RULE),
        ("BOX", (0, 0), (-1, -1), 0.6, RULE),
    ]
    if head_rows:
        style_cmds.append(("BACKGROUND", (0, 0), (-1, len(head_rows) - 1), TABLE_HEADER_BG))
        style_cmds.append(("LINEBELOW", (0, len(head_rows) - 1), (-1, len(head_rows) - 1), 0.8, MUTED))
    for index in range(len(head_rows), len(rows)):
        if (index - len(head_rows)) % 2 == 1:
            style_cmds.append(("BACKGROUND", (0, index), (-1, index), TABLE_ALT_BG))
    table.setStyle(TableStyle(style_cmds))
    return table


def _list_items(list_node: Tag, styles: dict[str, ParagraphStyle], ordered: bool) -> list:
    flowables = []
    for index, li in enumerate(list_node.find_all("li", recursive=False), 1):
        body_html = "".join(_inline_to_paragraph_html(child) for child in li.children if not (isinstance(child, Tag) and child.name in {"ul", "ol"}))
        bullet = f"{index}." if ordered else "•"
        para_style = styles["numbered"] if ordered else styles["bullet"]
        flowables.append(Paragraph(f"<font color='#2563eb'><b>{bullet}</b></font>&nbsp;&nbsp;{body_html.strip() or '&nbsp;'}", para_style))
        for sub in li.find_all(["ul", "ol"], recursive=False):
            sub_ordered = sub.name == "ol"
            for sub_flow in _list_items(sub, styles, sub_ordered):
                # Indent nested items.
                if isinstance(sub_flow, Paragraph):
                    sub_flow.style = ParagraphStyle(
                        f"{sub_flow.style.name}-nested",
                        parent=sub_flow.style,
                        leftIndent=sub_flow.style.leftIndent + 16,
                    )
                flowables.append(sub_flow)
    return flowables


def _build_flowables(soup: BeautifulSoup, styles: dict[str, ParagraphStyle]) -> list:
    flowables: list = []
    heading_styles = {"h1": "h1", "h2": "h2", "h3": "h3", "h4": "h4", "h5": "h4", "h6": "h4"}

    for node in soup.children:
        if isinstance(node, NavigableString):
            text = _strip(str(node))
            if text:
                flowables.append(Paragraph(text, styles["body"]))
            continue
        if not isinstance(node, Tag):
            continue
        tag = node.name.lower()
        if tag in heading_styles:
            inline = "".join(_inline_to_paragraph_html(child) for child in node.children) or "&nbsp;"
            flowables.append(Paragraph(inline, styles[heading_styles[tag]]))
            if tag == "h1":
                flowables.append(HRFlowable(width="100%", thickness=0.6, color=PRIMARY, spaceAfter=8))
        elif tag == "p":
            inline = "".join(_inline_to_paragraph_html(child) for child in node.children).strip()
            if inline:
                flowables.append(Paragraph(inline, styles["body"]))
        elif tag in {"ul", "ol"}:
            flowables.extend(_list_items(node, styles, ordered=(tag == "ol")))
            flowables.append(Spacer(1, 4))
        elif tag == "blockquote":
            inner_html = "".join(_inline_to_paragraph_html(child) for child in node.children).strip()
            inner_html = inner_html.replace("<br/><br/>", "<br/>")
            quote_style = ParagraphStyle(
                "QuoteRun",
                parent=styles["quote"],
                leftIndent=14,
                borderColor=QUOTE_BAR,
                borderWidth=2,
                borderPadding=(8, 12, 8, 12),
            )
            flowables.append(Paragraph(inner_html or "&nbsp;", quote_style))
        elif tag in {"pre"}:
            code_node = node.find("code") or node
            code_text = code_node.get_text("\n").rstrip()
            safe = (
                code_text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )
            flowables.append(Paragraph(safe or "&nbsp;", styles["code"]))
        elif tag == "table":
            flowables.append(Spacer(1, 4))
            flowables.append(_build_table(node, styles))
            flowables.append(Spacer(1, 8))
        elif tag == "hr":
            flowables.append(HRFlowable(width="100%", thickness=0.4, color=RULE, spaceBefore=6, spaceAfter=6))
        else:
            inline = "".join(_inline_to_paragraph_html(child) for child in node.children).strip()
            if inline:
                flowables.append(Paragraph(inline, styles["body"]))
    return flowables


def _draw_chrome(canvas, doc, *, title: str, generated_at: str) -> None:
    """Draw the page header and footer with title + page numbers."""
    canvas.saveState()
    body_font, bold_font, *_ = _register_fonts()
    width, height = A4

    canvas.setFillColor(MUTED)
    canvas.setFont(body_font, 8.5)
    canvas.drawString(2 * cm, height - 1.2 * cm, title)
    canvas.drawRightString(width - 2 * cm, height - 1.2 * cm, generated_at)
    canvas.setStrokeColor(RULE)
    canvas.setLineWidth(0.4)
    canvas.line(2 * cm, height - 1.35 * cm, width - 2 * cm, height - 1.35 * cm)

    canvas.setFillColor(MUTED)
    canvas.setFont(body_font, 8.5)
    page = f"Page {doc.page}"
    canvas.drawCentredString(width / 2, 1.3 * cm, page)
    canvas.setFont(body_font, 7.5)
    canvas.drawString(2 * cm, 1.3 * cm, "M1 Informatique · UVSQ / Université Paris-Saclay")
    canvas.drawRightString(width - 2 * cm, 1.3 * cm, "Assistant M1 AMIS")
    canvas.restoreState()


def render_pdf(title: str, body: str, *, generated_at: str | None = None) -> bytes:
    soup = _markdown_to_soup(body)
    styles = _build_styles()
    generated_at = generated_at or datetime.now().strftime("%d/%m/%Y · %H:%M")
    body_font, bold_font, italic_font, _ = _register_fonts()

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        title=title,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2.2 * cm,
        bottomMargin=2 * cm,
    )

    story: list = [
        Paragraph(title, styles["title"]),
        Paragraph(f"Généré le {generated_at} · M1 Informatique UVSQ / Université Paris-Saclay", styles["subtitle"]),
        HRFlowable(width="100%", thickness=0.6, color=PRIMARY, spaceAfter=10),
    ]
    story.extend(_build_flowables(soup, styles))

    on_page = lambda canvas, document: _draw_chrome(canvas, document, title=title, generated_at=generated_at)
    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# DOCX renderer
# --------------------------------------------------------------------------- #


def _docx_runs_from_inline(paragraph, node: Tag | NavigableString, *, bold=False, italic=False, mono=False, link: str | None = None) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        if link:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            run.font.underline = True
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    new_bold = bold or name in {"b", "strong"}
    new_italic = italic or name in {"i", "em"}
    new_mono = mono or name == "code"
    new_link = link or (node.get("href") if name == "a" else None)
    if name == "br":
        paragraph.add_run().add_break()
        return
    for child in node.children:
        _docx_runs_from_inline(paragraph, child, bold=new_bold, italic=new_italic, mono=new_mono, link=new_link)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _build_docx_table(document: Document, table_node: Tag) -> None:
    rows = table_node.find_all("tr")
    if not rows:
        return
    header_cells = rows[0].find_all(["th", "td"])
    column_count = len(header_cells)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    for row_index, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for col_index in range(column_count):
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""  # reset
            paragraph = cell.paragraphs[0]
            content = cells[col_index] if col_index < len(cells) else None
            if content is None:
                continue
            for child in content.children:
                _docx_runs_from_inline(paragraph, child, bold=(row_index == 0))
            if row_index == 0:
                _set_cell_shading(cell, "F1F5F9")


def _build_docx(document: Document, soup: BeautifulSoup) -> None:
    heading_levels = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

    for node in soup.children:
        if isinstance(node, NavigableString):
            text = _strip(str(node))
            if text:
                document.add_paragraph(text)
            continue
        if not isinstance(node, Tag):
            continue
        tag = node.name.lower()
        if tag in heading_levels:
            heading = document.add_heading("", level=heading_levels[tag])
            for child in node.children:
                _docx_runs_from_inline(heading, child)
        elif tag == "p":
            paragraph = document.add_paragraph()
            for child in node.children:
                _docx_runs_from_inline(paragraph, child)
        elif tag in {"ul", "ol"}:
            style = "List Number" if tag == "ol" else "List Bullet"
            for li in node.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style=style)
                for child in li.children:
                    if isinstance(child, Tag) and child.name in {"ul", "ol"}:
                        continue
                    _docx_runs_from_inline(paragraph, child)
        elif tag == "blockquote":
            paragraph = document.add_paragraph(style="Intense Quote")
            for child in node.children:
                _docx_runs_from_inline(paragraph, child)
        elif tag == "pre":
            code_node = node.find("code") or node
            paragraph = document.add_paragraph()
            run = paragraph.add_run(code_node.get_text("\n").rstrip())
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif tag == "table":
            _build_docx_table(document, node)
        elif tag == "hr":
            paragraph = document.add_paragraph("")
            run = paragraph.add_run("―" * 40)
            run.font.color.rgb = RGBColor(0xC8, 0xCC, 0xD3)
        else:
            paragraph = document.add_paragraph()
            for child in node.children:
                _docx_runs_from_inline(paragraph, child)


def render_docx(title: str, body: str, *, generated_at: str | None = None) -> bytes:
    soup = _markdown_to_soup(body)
    document = Document()

    style_normal = document.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    title_paragraph = document.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run = subtitle.add_run(
        f"Généré le {generated_at or datetime.now().strftime('%d/%m/%Y à %H:%M')} · "
        "M1 Informatique UVSQ / Université Paris-Saclay"
    )
    sub_run.italic = True
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    sub_run.font.size = Pt(10)

    _build_docx(document, soup)

    sections = document.sections
    for section in sections:
        footer = section.footer
        if footer.paragraphs:
            footer_paragraph = footer.paragraphs[0]
        else:
            footer_paragraph = footer.add_paragraph()
        footer_paragraph.text = "M1 Informatique · Assistant M1 AMIS"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# Public surface (preserves the old function names so callers don't break)
# --------------------------------------------------------------------------- #


def export_markdown_to_pdf(title: str, body: str) -> bytes:
    return render_pdf(title, body)


def export_markdown_to_docx(title: str, body: str) -> bytes:
    return render_docx(title, body)


__all__ = ["render_pdf", "render_docx", "export_markdown_to_pdf", "export_markdown_to_docx"]
